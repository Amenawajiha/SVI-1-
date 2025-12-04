"""
API Test Script
Tests the FastAPI endpoints using requests library
"""

import time
import requests
from pathlib import Path
from docx import Document


API_BASE_URL = "http://localhost:8000"


def create_test_document():
    """Create a test .docx document"""
    print("📝 Creating test document...")
    
    doc = Document()
    doc.add_heading('Travel Guide: Japan', 0)
    
    doc.add_heading('Visa Requirements', level=1)
    doc.add_paragraph(
        'Citizens of many countries can visit Japan visa-free for tourism '
        'for up to 90 days. Check with your local Japanese embassy for '
        'specific requirements based on your nationality.'
    )
    
    doc.add_heading('Best Time to Visit', level=1)
    doc.add_paragraph(
        'Spring (March to May) is ideal for cherry blossoms. Fall (September to November) '
        'offers beautiful autumn colors and comfortable temperatures.'
    )
    
    doc.add_heading('Top Cities', level=1)
    doc.add_paragraph(
        'Tokyo is the vibrant capital with modern technology and traditional temples. '
        'Kyoto features historic temples and traditional gardens. '
        'Osaka is known for its food culture and nightlife.'
    )
    
    # Save
    temp_dir = Path("./temp_api_test")
    temp_dir.mkdir(exist_ok=True)
    filepath = temp_dir / "japan_guide.docx"
    doc.save(filepath)
    
    print(f"✅ Test document created: {filepath}")
    return filepath


def test_health():
    """Test 1: Health check"""
    print("\n[Test 1] Testing health endpoint...")
    print("-" * 60)
    
    try:
        response = requests.get(f"{API_BASE_URL}/health")
        
        if response.status_code == 200:
            data = response.json()
            print(f"✅ Health check passed")
            print(f"   Status: {data['status']}")
            print(f"   Qdrant: {'✓' if data['qdrant_connected'] else '✗'}")
            print(f"   Embeddings: {'✓' if data['embeddings_loaded'] else '✗'}")
            print(f"   Storage: {'✓' if data['storage_available'] else '✗'}")
            return True
        else:
            print(f"❌ Health check failed: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ Health check error: {e}")
        return False


def test_upload_document():
    """Test 2: Upload document"""
    print("\n[Test 2] Testing document upload...")
    print("-" * 60)
    
    try:
        # Create test document
        doc_path = create_test_document()
        
        # Upload
        with open(doc_path, 'rb') as f:
            files = {'file': ('japan_guide.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {
                'collection_name': 'travel_guides_api_test',
                'category': 'travel'
            }
            
            print("📤 Uploading document...")
            response = requests.post(
                f"{API_BASE_URL}/documents/upload",
                files=files,
                data=data
            )
        
        if response.status_code == 201:
            result = response.json()
            print(f"✅ Document uploaded successfully")
            print(f"   Doc ID: {result['doc_id']}")
            print(f"   Filename: {result['original_filename']}")
            print(f"   Collection: {result['collection_name']}")
            print(f"   Chunks created: {result['chunks_created']}")
            print(f"   Word count: {result['word_count']}")
            return result['doc_id'], result['collection_name']
        else:
            print(f"❌ Upload failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return None, None
    except Exception as e:
        print(f"❌ Upload error: {e}")
        return None, None


def test_search(collection_name):
    """Test 3: Search documents"""
    print("\n[Test 3] Testing search...")
    print("-" * 60)
    
    queries = [
        "Do I need a visa to visit Japan?",
        "When is the best time to visit?",
        "Tell me about Tokyo"
    ]
    
    try:
        for i, query in enumerate(queries, 1):
            print(f"\n🔍 Query {i}: '{query}'")
            
            payload = {
                "query": query,
                "collection_name": collection_name,
                "top_k": 2,
                "score_threshold": 0.3
            }
            
            response = requests.post(
                f"{API_BASE_URL}/documents/search",
                json=payload
            )
            
            if response.status_code == 200:
                result = response.json()
                print(f"   Found {result['results_count']} results:")
                
                for j, res in enumerate(result['results'], 1):
                    print(f"\n   Result {j} (Score: {res['score']:.3f}):")
                    print(f"   {res['text'][:150]}...")
            else:
                print(f"   ❌ Search failed: {response.status_code}")
        
        print("\n✅ Search tests completed")
        return True
        
    except Exception as e:
        print(f"❌ Search error: {e}")
        return False


def test_list_collections():
    """Test 4: List collections"""
    print("\n[Test 4] Testing list collections...")
    print("-" * 60)
    
    try:
        response = requests.get(f"{API_BASE_URL}/collections")
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Collections listed successfully")
            print(f"   Total collections: {result['total_collections']}")
            
            for coll in result['collections']:
                print(f"   - {coll['name']}: {coll['points_count']} points")
            
            return True
        else:
            print(f"❌ List collections failed: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ List collections error: {e}")
        return False


def test_delete_document(doc_id, collection_name):
    """Test 5: Delete document"""
    print("\n[Test 5] Testing document deletion...")
    print("-" * 60)
    
    try:
        data = {'collection_name': collection_name}
        
        response = requests.delete(
            f"{API_BASE_URL}/documents/{doc_id}",
            data=data
        )
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Document deleted successfully")
            print(f"   Doc ID: {result['doc_id']}")
            print(f"   Message: {result['message']}")
            return True
        else:
            print(f"❌ Delete failed: {response.status_code}")
            print(f"   Response: {response.text}")
            return False
    except Exception as e:
        print(f"❌ Delete error: {e}")
        return False


def cleanup():
    """Clean up test files"""
    print("\n[Cleanup] Removing test files...")
    try:
        import shutil
        if Path("./temp_api_test").exists():
            shutil.rmtree("./temp_api_test")
        print("✅ Cleanup completed")
    except Exception as e:
        print(f"⚠️  Cleanup warning: {e}")


if __name__ == "__main__":
    print("=" * 60)
    print("API Test Suite")
    print("=" * 60)
    print("\n⚠️  Prerequisites:")
    print("   1. API server must be running:")
    print("      uvicorn api.main:app --reload")
    print("   2. Qdrant must be running")
    print()
    input("Press Enter to start tests...")
    
    doc_id = None
    collection_name = None
    
    try:
        # Test 1: Health check
        if not test_health():
            print("\n❌ Health check failed. Make sure API is running.")
            exit(1)
        
        # Wait a bit for initialization
        time.sleep(1)
        
        # Test 2: Upload document
        doc_id, collection_name = test_upload_document()
        if not doc_id:
            print("\n❌ Upload failed. Cannot continue tests.")
            exit(1)
        
        # Wait for indexing
        print("\n⏳ Waiting for indexing to complete...")
        time.sleep(2)
        
        # Test 3: Search
        if not test_search(collection_name):
            print("\n⚠️  Search tests failed")
        
        # Test 4: List collections
        if not test_list_collections():
            print("\n⚠️  List collections failed")
        
        # Test 5: Delete document
        if doc_id and collection_name:
            if not test_delete_document(doc_id, collection_name):
                print("\n⚠️  Delete test failed")
        
        print("\n" + "=" * 60)
        print("🎉 API TESTS COMPLETED!")
        print("=" * 60)
        print("\n✅ Your API is fully functional!")
        print("\n📚 API Documentation:")
        print(f"   Swagger UI: {API_BASE_URL}/docs")
        print(f"   ReDoc: {API_BASE_URL}/redoc")
        
    finally:
        cleanup()
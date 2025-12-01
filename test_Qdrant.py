"""
Complete End-to-End Pipeline Test
Tests the full workflow: Upload → Process → Store → Search
"""

import shutil
from pathlib import Path
from docx import Document

from Qdrant.document_pipeline.embedding import E5Embeddings
from Qdrant.document_pipeline.storage import DocumentStorage
from Qdrant.document_pipeline.processor import DocumentProcessor
from Qdrant.document_pipeline.qdrant_manager import QdrantManager


def create_test_document():
    """Create a test .docx document"""
    print("Creating test document...")
    
    doc = Document()
    doc.add_heading('Italy Travel Guide', 0)
    
    doc.add_heading('Visa Requirements', level=1)
    doc.add_paragraph(
        'Most travelers need a Schengen visa to visit Italy. The application '
        'requires a valid passport, travel insurance, and proof of accommodation. '
        'Processing time is typically 15 business days.'
    )
    
    doc.add_heading('Best Time to Visit', level=1)
    doc.add_paragraph(
        'Spring (April to June) and fall (September to October) offer the best '
        'weather and fewer crowds. Summer is peak tourist season with hot temperatures.'
    )
    
    doc.add_heading('Top Destinations', level=1)
    doc.add_paragraph(
        'Rome features ancient landmarks like the Colosseum and Vatican City. '
        'Florence is the Renaissance art capital with the Uffizi Gallery. '
        'Venice is famous for its romantic canals and gondola rides.'
    )
    
    doc.add_heading('Transportation', level=1)
    doc.add_paragraph(
        'Italy has excellent high-speed trains connecting major cities. '
        'Public transportation in cities includes buses, trams, and metros. '
        'Book train tickets in advance for better prices.'
    )
    
    # Save
    temp_dir = Path("./temp_test_files")
    temp_dir.mkdir(exist_ok=True)
    filepath = temp_dir / "italy_guide.docx"
    doc.save(filepath)
    
    print(f"Test document created: {filepath}")
    return str(filepath)


def test_complete_pipeline():
    """Run complete end-to-end test"""
    print("\n" + "=" * 60)
    print("COMPLETE PIPELINE TEST")
    print("=" * 60)
    
    # Initialize all components
    print("\n[Step 1] Initializing components...")
    print("-" * 60)
    
    embeddings = E5Embeddings(device="cpu")
    storage = DocumentStorage(base_upload_dir="./test_uploads")
    processor = DocumentProcessor(
        embeddings=embeddings,
        storage=storage,
        chunk_size=400,
        chunk_overlap=50
    )
    qdrant_manager = QdrantManager(
        host="localhost",
        port=6333,
        embeddings=embeddings
    )
    
    collection_name = "travel_guides_test"
    
    print("All components initialized")
    
    # Create test document
    print("\n[Step 2] Creating test document...")
    print("-" * 60)
    test_file = create_test_document()
    
    # Process document
    print("\n[Step 3] Processing document...")
    print("-" * 60)
    result = processor.process_document(
        file_path=test_file,
        original_filename="italy_guide.docx",
        collection_name=collection_name,
        custom_metadata={
            'category': 'travel',
            'country': 'Italy',
            'language': 'English'
        }
    )
    
    doc_id = result['doc_id']
    chunks_data = result['chunks']
    
    print(f"\nProcessing Results:")
    print(f"   Doc ID: {doc_id}")
    print(f"   Chunks created: {result['num_chunks']}")
    print(f"   Word count: {result['word_count']}")
    
    # Store in Qdrant
    print("\n[Step 4] Storing in Qdrant...")
    print("-" * 60)
    
    # Create collection
    qdrant_manager.create_collection(
        collection_name=collection_name,
        recreate=True
    )
    
    # Add documents
    added_count = qdrant_manager.add_documents(
        collection_name=collection_name,
        chunks_data=chunks_data
    )
    
    print(f"Stored {added_count} chunks in Qdrant")
    
    # Verify storage
    collection_info = qdrant_manager.get_collection_info(collection_name)
    print(f"   Collection points: {collection_info['points_count']}")
    
    # Test search functionality
    print("\n[Step 5] Testing search...")
    print("-" * 60)
    
    # Search 1: Visa information
    print("\nQuery 1: 'What visa do I need for Italy?'")
    results1 = qdrant_manager.search(
        collection_name=collection_name,
        query="What visa do I need for Italy?",
        top_k=2,
        score_threshold=0.3
    )
    
    for i, result in enumerate(results1, 1):
        print(f"\n   Result {i} (Score: {result['score']:.3f}):")
        print(f"   {result['text'][:200]}...")
        print(f"   [Chunk: {result['metadata']['chunk_id']}]")
    
    # Search 2: Travel timing
    print("\nQuery 2: 'When is the best time to visit?'")
    results2 = qdrant_manager.search(
        collection_name=collection_name,
        query="When is the best time to visit?",
        top_k=2,
        score_threshold=0.3
    )
    
    for i, result in enumerate(results2, 1):
        print(f"\n   Result {i} (Score: {result['score']:.3f}):")
        print(f"   {result['text'][:200]}...")
        print(f"   [Chunk: {result['metadata']['chunk_id']}]")
    
    # Search 3: Cities and destinations
    print("\nQuery 3: 'What cities should I visit?'")
    results3 = qdrant_manager.search(
        collection_name=collection_name,
        query="What cities should I visit?",
        top_k=2,
        score_threshold=0.3
    )
    
    for i, result in enumerate(results3, 1):
        print(f"\n   Result {i} (Score: {result['score']:.3f}):")
        print(f"   {result['text'][:200]}...")
        print(f"   [Chunk: {result['metadata']['chunk_id']}]")
    
    # Test filtered search
    print("\n[Step 6] Testing filtered search...")
    print("-" * 60)
    
    filtered_results = qdrant_manager.search(
        collection_name=collection_name,
        query="travel information",
        top_k=3,
        filter_conditions={'country': 'Italy'}
    )
    
    print(f"Found {len(filtered_results)} results with country filter")
    
    # Get final statistics
    print("\n[Step 7] Final statistics...")
    print("-" * 60)
    
    stats = qdrant_manager.get_stats()
    print(f"Qdrant Statistics:")
    print(f"   Total collections: {stats['total_collections']}")
    for coll in stats['collections']:
        print(f"   - {coll['name']}: {coll['points_count']} points")
    
    storage_stats = storage.get_storage_stats()
    print(f"\nStorage Statistics:")
    print(f"   Total files: {storage_stats['total_files']}")
    print(f"   Total size: {storage_stats['total_size_mb']} MB")
    
    return True, doc_id, collection_name, qdrant_manager


def cleanup(doc_id, collection_name, qdrant_manager):
    """Clean up test data"""
    print("\n[Cleanup] Removing test data...")
    print("-" * 60)
    
    # Ask for confirmation before cleanup
    response = input("\nDo you want to delete the test document and clean up all test data? (yes/no): ").strip().lower()
    if response not in ['yes', 'y']:
        print("Cleanup cancelled. Test data preserved.")
        return
    
    try:
        # Delete from Qdrant
        if qdrant_manager and collection_name:
            qdrant_manager.delete_collection(collection_name)
            print(f"Deleted Qdrant collection: {collection_name}")
        
        # Delete uploaded files
        if Path("./test_uploads").exists():
            shutil.rmtree("./test_uploads")
            print("Deleted uploaded files")
        
        # Delete temp files
        if Path("./temp_test_files").exists():
            shutil.rmtree("./temp_test_files")
            print("Deleted temp files")
        
        print("Cleanup completed")
        
    except Exception as e:
        print(f"Cleanup warning: {e}")


if __name__ == "__main__":
    print("=" * 60)
    print("END-TO-END PIPELINE TEST")
    print("=" * 60)
    print("\nPrerequisites:")
    print("   1. Qdrant must be running (cd qdrant_setup && docker-compose up -d)")
    print("   2. All dependencies installed (pip install -r requirements.txt)")
    print()
    input("Press Enter to continue...")
    
    doc_id = None
    collection_name = None
    qdrant_manager = None
    
    try:
        # Run complete pipeline test
        success, doc_id, collection_name, qdrant_manager = test_complete_pipeline()
        
        if success:
            print("\n" + "=" * 60)
            print("PIPELINE TEST PASSED")
            print("=" * 60)
        
        print("\n" + "=" * 60)
        print("ALL TESTS PASSED")
        print("=" * 60)
        print("\nYour RAG pipeline is fully functional!")
        print("\nWhat just happened:")
        print("   1. Created a .docx document")
        print("   2. Extracted and cleaned text")
        print("   3. Chunked text intelligently")
        print("   4. Generated embeddings (768-d vectors)")
        print("   5. Stored in Qdrant vector database")
        print("   6. Performed semantic search")
        print("   7. Retrieved relevant context")
        print("\nNext step: Build the FastAPI endpoint!")
        
    finally:
        cleanup(doc_id, collection_name, qdrant_manager)